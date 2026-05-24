"""
Automatyzator Informacji Dodatkowej do sprawozdania finansowego
Streamlit app z Claude 3.5 Sonnet + LlamaParse + python-docx
"""

import streamlit as st
import anthropic
import os
import io
import json
import re
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import requests
from datetime import date, datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

# ─── PAGE CONFIG ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Generator Informacji Dodatkowej",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #1e3a5f 0%, #2d6a9f 100%);
        color: white; padding: 2rem; border-radius: 12px;
        margin-bottom: 2rem; text-align: center;
    }
    .step-card {
        background: #f8f9fa; border-left: 4px solid #2d6a9f;
        padding: 1rem 1.5rem; border-radius: 0 8px 8px 0; margin: 0.5rem 0;
    }
    .validation-ok  { color: #28a745; font-weight: bold; }
    .validation-warn{ color: #ffc107; font-weight: bold; }
    .validation-err { color: #dc3545; font-weight: bold; }
    .metric-box {
        background: white; border: 1px solid #dee2e6; border-radius: 8px;
        padding: 1rem; text-align: center; box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .stProgress > div > div { background-color: #2d6a9f; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h1>📊 Generator Informacji Dodatkowej</h1>
    <p>Automatyczne tworzenie not do sprawozdania finansowego zgodnie z Ustawą o Rachunkowości</p>
</div>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# MODUŁ 1: PARSOWANIE PDF
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf_basic(pdf_bytes: bytes, filename: str) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        text_parts = [f"=== DOKUMENT: {filename} ===\n"]
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"\n--- Strona {i+1} ---\n{page_text}")
        return "\n".join(text_parts)
    except Exception as e:
        return f"[BŁĄD ekstrakcji {filename}: {e}]"


def parse_documents_llamaparse(pdf_files, llama_api_key, progress_callback=None):
    try:
        from llama_parse import LlamaParse
        parser = LlamaParse(
            api_key=llama_api_key, result_type="markdown", language="pl",
            parsing_instruction=(
                "Dokument to sprawozdanie finansowe polskiej spółki. "
                "Zachowaj strukturę tabel finansowych."
            )
        )
        results = {}
        for idx, f in enumerate(pdf_files):
            if progress_callback:
                progress_callback(idx / len(pdf_files), f"Parsowanie: {f.name}")
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(f.getvalue()); tmp_path = tmp.name
            try:
                docs = parser.load_data(tmp_path)
                results[f.name] = "\n\n".join(d.text for d in docs)
            finally:
                os.unlink(tmp_path)
        return results
    except Exception as e:
        st.warning(f"⚠️ LlamaParse niedostępny ({e}) — używam pypdf.")
        return parse_documents_fallback(pdf_files, progress_callback)


def parse_documents_fallback(pdf_files, progress_callback=None):
    results = {}
    for idx, f in enumerate(pdf_files):
        if progress_callback:
            progress_callback(idx / len(pdf_files), f"Ekstrakcja: {f.name}")
        results[f.name] = extract_text_from_pdf_basic(f.getvalue(), f.name)
    return results


# ═══════════════════════════════════════════════════════════════════════════════
# MODUŁ KRS
# ═══════════════════════════════════════════════════════════════════════════════

def fetch_krs_by_krs_nr(krs_nr: str):
    krs_clean = re.sub(r"[^0-9]", "", krs_nr).zfill(10)
    if len(krs_clean) != 10:
        return None
    headers = {"Accept": "application/json",
                "User-Agent": "Mozilla/5.0 (compatible; InformacjaDodatkowa/1.0)"}
    try:
        url = f"https://api-krs.ms.gov.pl/api/krs/OdpisAktualny/{krs_clean}"
        for rejestr in ["P", "S"]:
            r = requests.get(url, params={"rejestr": rejestr, "format": "json"},
                             headers=headers, timeout=20)
            if r.status_code == 200:
                return _parse_odpis(r.json(), krs_clean)
    except requests.exceptions.ConnectionError:
        raise ConnectionError("Brak połączenia z API KRS")
    except requests.exceptions.Timeout:
        raise TimeoutError("API KRS nie odpowiada")
    except Exception as e:
        raise RuntimeError(f"Błąd API KRS: {e}")
    return None


def fetch_krs_by_krs_nr_debug(krs_nr: str):
    import json as _json
    krs_clean = re.sub(r"[^0-9]", "", krs_nr).zfill(10)
    log = [f"Numer KRS po oczyszczeniu: {krs_clean}"]
    headers = {"Accept": "application/json",
                "User-Agent": "Mozilla/5.0 (compatible; InformacjaDodatkowa/1.0)"}
    url = f"https://api-krs.ms.gov.pl/api/krs/OdpisAktualny/{krs_clean}"
    for rejestr in ["P", "S"]:
        try:
            r = requests.get(url, params={"rejestr": rejestr, "format": "json"},
                             headers=headers, timeout=20)
            log.append(f"\n→ {url}?rejestr={rejestr}")
            log.append(f"  Status: {r.status_code}")
            if r.status_code == 200:
                data = r.json()
                log.append(f"  Odpowiedź:\n{_json.dumps(data, ensure_ascii=False, indent=2)[:2000]}")
                parsed = _parse_odpis(data, krs_clean)
                if parsed:
                    log.append("\n✅ Dane sparsowane pomyślnie")
                    return parsed, "\n".join(log)
            else:
                log.append(f"  Błąd: {r.text[:200]}")
        except Exception as e:
            log.append(f"  Wyjątek: {e}")
    log.append("\n❌ Nie udało się pobrać danych.")
    return None, "\n".join(log)


def _parse_odpis(data: dict, krs_nr: str = ""):
    try:
        odpis   = data.get("odpis", data)
        naglowek = odpis.get("naglowekA", {})
        dane    = odpis.get("dane", {})
        dzial1  = dane.get("dzial1", {})
        dane_p  = dzial1.get("danePodmiotu", {})
        ident   = dane_p.get("identyfikatory", {})
        siedz   = dzial1.get("siedzibaIAdres", {}).get("adres", {})

        ulica   = siedz.get("ulica", "").replace("UL. ", "ul. ").replace("UL.", "ul.")
        nr_domu = siedz.get("nrDomu", "")
        nr_lok  = siedz.get("nrLokalu", "")
        kod     = siedz.get("kodPocztowy", "")
        miasto  = siedz.get("miejscowosc", "")
        siedziba = f"{ulica} {nr_domu}".strip()
        if nr_lok:
            siedziba += f"/{nr_lok}"
        if kod and miasto:
            siedziba += f", {kod} {miasto}"

        pkd = ""
        def _wyciagnij_pkd(s):
            lista = s.get("przedmiotPrzewazajacejDzialalnosci") or s.get("przedmiotDzialalnosci") or []
            if isinstance(lista, list) and lista:
                p = lista[0]
                return f"{p.get('kodDzialalnosci','')} {p.get('opis','')}".strip()
            if isinstance(lista, dict):
                return f"{lista.get('kodDzialalnosci','')} {lista.get('opis','')}".strip()
            return ""

        for k in ["dzial1", "dzial3", "dzial2"]:
            sec = dane.get(k, {}).get("przedmiotDzialalnosci", {})
            if sec:
                pkd = _wyciagnij_pkd(sec)
                if pkd:
                    break

        regon_raw = ident.get("regon", "")
        return {
            "nazwa":           dane_p.get("nazwa", ""),
            "siedziba":        siedziba,
            "nip":             ident.get("nip", ""),
            "krs":             naglowek.get("numerKRS", krs_nr),
            "regon":           regon_raw[:9] if regon_raw else "",
            "pkd":             pkd,
            "data_rejestracji": naglowek.get("dataRejestracjiWKRS", ""),
            "forma_prawna":    dane_p.get("formaPrawna", ""),
        }
    except Exception:
        return None


# ═══════════════════════════════════════════════════════════════════════════════
# MODUŁ 2: MAPOWANIE DOKUMENTÓW
# ═══════════════════════════════════════════════════════════════════════════════

REQUIRED_DOC_TYPES = {
    "BILANS": {
        "label": "Bilans", "icon": "🏦",
        "desc": "Zestawienie aktywów i pasywów na dzień bilansowy",
        "keywords": ["aktywa trwałe", "aktywa obrotowe", "pasywa", "kapitał własny", "zobowiązania"],
    },
    "RZiS": {
        "label": "Rachunek Zysków i Strat", "icon": "📈",
        "desc": "Przychody, koszty i wynik finansowy za rok obrotowy",
        "keywords": ["przychody ze sprzedaży", "koszty działalności", "zysk netto", "wynik finansowy"],
    },
    "ŚRODKI TRWAŁE": {
        "label": "Tabela środków trwałych", "icon": "🏗️",
        "desc": "Wartość brutto, umorzenia i wartość netto środków trwałych",
        "keywords": ["środki trwałe", "wartość brutto", "umorzenie", "odpisy amortyzacyjne"],
    },
    "PRZEPŁYWY PIENIĘŻNE": {
        "label": "Rachunek przepływów pieniężnych", "icon": "💸",
        "desc": "Cash flow: operacyjny, inwestycyjny, finansowy",
        "keywords": ["przepływy", "działalność operacyjna", "działalność inwestycyjna"],
    },
    "POLITYKA RACHUNKOWOŚCI": {
        "label": "Polityka rachunkowości", "icon": "📜",
        "desc": "Przyjęte zasady rachunkowości, metody wyceny, okresy amortyzacji",
        "keywords": ["polityka rachunkowości", "zasady rachunkowości", "metody wyceny"],
    },
    "ZOiS": {
        "label": "Zestawienie Obrotów i Sald", "icon": "⚖️",
        "desc": "Obroty i salda kont księgi głównej za rok obrotowy",
        "keywords": ["zestawienie obrotów", "obroty i salda", "salda końcowe",
                     "konta aktywne", "obroty debetowe", "saldo dt", "saldo ct",
                     "zestawienie kont", "plan kont", "księga główna"],
    },
}

HEADER_RULES = [
    ("ZOiS",                   ["zestawienie obrotów i sald", "obroty i salda", "zois"]),
    ("BILANS",                  ["bilans na dzień", "bilans jednostki", "aktywa i pasywa"]),
    ("RZiS",                    ["rachunek zysków i strat", "rachunek wyników", "wynik finansowy netto"]),
    ("ŚRODKI TRWAŁE",           ["tabela środków trwałych", "środki trwałe i wartości",
                                  "zestawienie środków trwałych"]),
    ("POLITYKA RACHUNKOWOŚCI",  ["polityka rachunkowości", "zasady rachunkowości przyjęte"]),
    ("PRZEPŁYWY PIENIĘŻNE",     ["rachunek przepływów pieniężnych", "cash flow"]),
]


def identify_document_type(text: str) -> str:
    header = text[:500].lower()
    for doc_type, phrases in HEADER_RULES:
        if any(p in header for p in phrases):
            return doc_type
    text_lower = text.lower()
    scores = {dt: sum(text_lower.count(kw) for kw in info["keywords"])
              for dt, info in REQUIRED_DOC_TYPES.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "INNY"


def check_missing_documents(doc_mapping: dict) -> list:
    types_found = {d["type"] for d in doc_mapping.values()}
    return [dt for dt in REQUIRED_DOC_TYPES if dt not in types_found]


def map_documents(parsed_docs: dict) -> dict:
    return {
        filename: {"type": identify_document_type(text), "text": text, "length": len(text)}
        for filename, text in parsed_docs.items()
    }


# ═══════════════════════════════════════════════════════════════════════════════
# MODUŁ 3: WALIDACJA
# ═══════════════════════════════════════════════════════════════════════════════

def extract_financial_number(text: str, pattern: str):
    try:
        matches = re.findall(rf"{pattern}[:\s]+([+-]?\d[\d\s.,]*)", text, re.IGNORECASE)
        if matches:
            return float(matches[0].replace(" ", "").replace(",", "."))
    except Exception:
        pass
    return None


def validate_data_consistency(doc_mapping: dict) -> list:
    issues = []
    all_text = "\n".join(d["text"] for d in doc_mapping.values())

    aktywne = extract_financial_number(all_text, r"AKTYWA\s+RAZEM|suma\s+aktywów")
    pasywa  = extract_financial_number(all_text, r"PASYWA\s+RAZEM|suma\s+pasywów")
    if aktywne and pasywa:
        diff = abs(aktywne - pasywa)
        if diff < 1:
            issues.append({"level": "OK",   "msg": f"✅ Bilans zbilansowany: {aktywne:,.0f} PLN"})
        elif diff < aktywne * 0.001:
            issues.append({"level": "WARN", "msg": f"⚠️ Drobna różnica bilansowa: {diff:,.2f} PLN"})
        else:
            issues.append({"level": "ERR",  "msg": f"❌ Bilans NIE jest zbilansowany! Różnica: {diff:,.0f} PLN"})
    else:
        issues.append({"level": "WARN", "msg": "⚠️ Nie znaleziono sum bilansowych"})

    types_found = [d["type"] for d in doc_mapping.values()]
    for dt, info in REQUIRED_DOC_TYPES.items():
        lvl = "OK" if dt in types_found else "WARN"
        msg = f"✅ Znaleziono: {info['icon']} {info['label']}" if dt in types_found \
              else f"⚠️ Brak: {info['icon']} {info['label']}"
        issues.append({"level": lvl, "msg": msg})
    return issues


# ═══════════════════════════════════════════════════════════════════════════════
# MODUŁ 4: GENEROWANIE PRZEZ CLAUDE
# ═══════════════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """Jesteś biegłym rewidentem i ekspertem ds. rachunkowości polskiej.
Sporządzasz "Informację Dodatkową" do sprawozdania finansowego zgodnie z UoR.

ZASADA NADRZĘDNA:
- Generuj TYLKO noty z niezerowymi wartościami wynikającymi z dokumentów
- Jeśli nota miałaby zawierać same zera lub nie ma danych — POMIJASZ ją całkowicie
- NIE piszesz tytułu pominiętej noty, NIE piszesz "nie dotyczy", po prostu jej nie ma
- Gdzie brakuje konkretnych danych szczegółowych (np. podział należności wg terminów) — napisz [DO UZUPEŁNIENIA]

═══════════════════════════════════════════════════════
CZĘŚĆ 1. WPROWADZENIE
═══════════════════════════════════════════════════════

1.1 Dane identyfikacyjne — na podstawie danych z panelu (nazwa, forma prawna, siedziba, NIP, KRS, REGON, PKD, data rejestracji, okres sprawozdawczy, kontynuacja działalności)
1.2 Zasady rachunkowości — podstawa prawna, wariant RZiS, rodzaj sprawozdania, leasing, podatek odroczony
1.3 Metody wyceny — WNiP, ST, zapasy (FIFO/LIFO/śr.ważona), należności, zobowiązania, waluty obce
1.4 Amortyzacja — metoda, stawki/okresy dla grup ST i WNiP, próg niskocennych
1.5 Przychody i koszty — moment ujęcia, rezerwy, RMK
1.6 Korekty i zmiany polityki — jeśli nie było: jedno zdanie

═══════════════════════════════════════════════════════
CZĘŚĆ 2. NOTY (tylko z niezerowymi wartościami)
═══════════════════════════════════════════════════════

Każda nota = tabela Markdown. Pomiń bez komentarza gdy wartości = 0.

AKTYWA TRWAŁE:
Nota 1  — ST: grupy × (brutto BO, zwiększenia, zmniejszenia, brutto BZ, umorzenie BO, amortyzacja, umorzenie BZ, netto BZ) → gdy ST > 0
Nota 2  — WNiP: analogiczna struktura → gdy WNiP > 0
Nota 3  — Inwestycje długoterminowe → gdy > 0
Nota 4  — Odpisy długoterminowych aktywów niefinansowych → gdy są
Nota 5  — Odpisy długoterminowych aktywów finansowych → gdy są
Nota 6  — Koszty prac rozwojowych, wartość firmy → gdy > 0
Nota 7  — Grunty wieczyste → gdy są
Nota 32 — Odpisy aktualizujące ST → gdy są
Nota 38 — Nakłady na niefinansowe aktywa trwałe → gdy są

NALEŻNOŚCI:
Nota 10 — Odpisy aktualizujące należności: grupy × (BO, zwiększenia, wykorzystanie, uznanie za zbędne, BZ) → gdy odpisy > 0
Nota 60 — Struktura należności (przeterminowane/nieprzeterminowane) → gdy należności > 0
Nota 61 — Należności wg terminów (do 30 dni, 31-90, 91-180, >181, >12 mies.) → gdy należności > 0

KAPITAŁY:
Nota 12 — Kapitał podstawowy sp. z o.o.: wspólnik, wartość udziałów, % → zawsze dla sp. z o.o.
Nota 11 — Kapitał podstawowy SA/PSA: serie akcji → dla SA i PSA
Nota 13 — Zmiany kapitału zapasowego i rezerwowego: BO, zwiększenia, zmniejszenia, BZ → gdy ≠ 0
Nota 14 — Kapitał z aktualizacji wyceny → gdy > 0
Nota 15 — Podział zysku: zysk netto, wynik lat ubiegłych, razem, proponowany podział → gdy zysk > 0
Nota 16 — Pokrycie straty → gdy strata

ZOBOWIĄZANIA I REZERWY:
Nota 17 — Rezerwy: rodzaj × (BO, zwiększenia, wykorzystanie, rozwiązanie, BZ) → gdy > 0
Nota 18 — Podatek odroczony: aktywa i rezerwy, różnice przejściowe → gdy ≠ 0
Nota 19 — Zobowiązania wg wymagalności: rodzaje × (do 1 roku, 1-3 lata, 3-5 lat, >5 lat) BO i BZ → gdy > 0
Nota 20 — Zobowiązania zabezpieczone na majątku → gdy są
Nota 25 — Zobowiązania warunkowe (poręczenia, gwarancje) → gdy są
Nota 45 — Zobowiązania emerytalne → gdy są
Nota 72 — Gwarancje i poręczenia udzielone → gdy są
Nota 73 — Zobowiązania DT > 5 lat → gdy są

ROZLICZENIA MIĘDZYOKRESOWE:
Nota 21 — RMK czynne: wyszczególnienie × (BO, zwiększenia, zmniejszenia, BZ) → gdy > 0
Nota 22 — RMK przychodów → gdy > 0

PRZYCHODY I KOSZTY:
Nota 29 — Struktura przychodów (kraj/eksport/WDT): wyroby/usługi/towary × rok poprz./bież. → gdy przychody > 0
Nota 31 — Koszty rodzajowe: amortyzacja, materiały, usługi obce, podatki, wynagrodzenia, ubezpieczenia, pozostałe, razem × (rok poprzedni, rok bieżący) → ZAWSZE gdy RZiS dostępny
Nota 34 — Działalność zaniechana → gdy była
Nota 39 — Pozycje nadzwyczajne → gdy są

PODATKI:
Nota 35 — Rozliczenie CIT: zysk brutto, różnice trwałe, różnice przejściowe, podstawa, podatek należny → gdy CIT > 0
Nota 59 — Zapłacony CIT: CIT z RZiS, zmiana rezerwy, CIT wg deklaracji, zmiana należności, CIT zapłacony → gdy CIT > 0

ŚRODKI PIENIĘŻNE:
Nota 41 — Struktura środków: kasa, rachunki × (rok poprz., rok bież., zmiana, ograniczone) → gdy > 0
Nota 63 — Rachunek VAT (split payment) → gdy są środki

ZATRUDNIENIE I WYNAGRODZENIA:
Nota 43 — Zatrudnienie: umysłowi, robotnicy, razem × (rok poprz., rok bież.) → gdy > 0
Nota 44 — Wynagrodzenia organów (zarząd, RN) → gdy były
Nota 46 — Zaliczki/pożyczki dla organów → gdy były

JEDNOSTKI POWIĄZANE:
Nota 52 — Zaangażowanie kapitałowe → gdy są udziały w innych spółkach
Nota 76 — Transakcje z powiązanymi: jednostka × (charakter, należności, zobowiązania, przychody, koszty) → gdy są

═══════════════════════════════════════════════════════
CZĘŚĆ 3. POZOSTAŁE INFORMACJE
═══════════════════════════════════════════════════════
3.1 Zdarzenia po dniu bilansowym
3.2 Inne istotne informacje

═══════════════════════════════════════════════════════
CZĘŚĆ 4. DANE DO WYKRESÓW (JSON — OBOWIĄZKOWE)
═══════════════════════════════════════════════════════

Na końcu odpowiedzi, po wszystkich notach, umieść DOKŁADNIE ten blok JSON
(wypełniony rzeczywistymi danymi z dokumentów):

```wykres_dane
{
  "wynik": {
    "tytul": "Wynik finansowy",
    "etykiety": ["Przychody", "Koszty operacyjne", "Zysk/Strata netto"],
    "rok_poprzedni": [0, 0, 0],
    "rok_biezacy": [0, 0, 0]
  },
  "koszty": {
    "tytul": "Struktura kosztów rodzajowych",
    "etykiety": ["Amortyzacja", "Materiały i energia", "Usługi obce", "Wynagrodzenia", "Ubezpieczenia", "Pozostałe"],
    "rok_poprzedni": [0, 0, 0, 0, 0, 0],
    "rok_biezacy": [0, 0, 0, 0, 0, 0]
  },
  "pasywa": {
    "tytul": "Struktura pasywów",
    "etykiety": ["Kapitał własny", "Zobowiązania długoterminowe", "Zobowiązania krótkoterminowe", "RMK"],
    "rok_biezacy": [0, 0, 0, 0]
  }
}
```

Zastąp zera rzeczywistymi wartościami z dokumentów. Pomiń pozycje z wartością 0 z etykiet.
Użyj liczb bez spacji i przecinków (np. 311208.30 nie 311 208,30).

FORMAT TREŚCI:
- Nagłówki Markdown (##, ###)
- Tabele Markdown dla not
- Liczby: 1 234 567,89 PLN (z separatorem tysięcy, przecinek dziesiętny)
- NIE pisz "Nie dotyczy", NIE generuj not z samymi zerami, NIE pisz tytułu pominiętej noty
"""


def generate_accounting_notes(doc_mapping: dict, anthropic_api_key: str,
                               company_name: str, year: int,
                               company_info: dict = None,
                               progress_callback=None) -> str:
    client = anthropic.Anthropic(api_key=anthropic_api_key)
    info = company_info or {}

    pa = info.get("polityka_answers", {})
    polityka_blok = ""
    if pa:
        polityka_blok = (
            "\n📋 ZASADY RACHUNKOWOŚCI (z ankiety — brak Polityki Rachunkowości):\n"
            f"- Wynik finansowy: {pa.get('wynik_finansowy','')}\n"
            f"- Wycena zapasów: {pa.get('wycena_zapasow','')}\n"
            f"- Amortyzacja ST: {pa.get('amortyzacja','')}\n"
            f"- Wycena należności: {pa.get('wycena_naleznosci','')}\n"
            f"- Rodzaj sprawozdania: {pa.get('sposob_sprawozdania','')}\n"
            f"- Podatek odroczony: {'TAK' if pa.get('podatek_odroczony') else 'NIE'}\n"
            f"- Leasing: {pa.get('leasing','')}\n"
            + (f"- Uwagi: {pa['uwagi']}\n" if pa.get("uwagi") else "")
            + "Wypełnij sekcje 1.2–1.5 na podstawie powyższych danych.\n"
        )

    zagrozenie_blok = ""
    if info.get("zagrozenie_kontynuacji"):
        zagrozenie_blok = (
            "\n⚠️ ZAGROŻENIE KONTYNUACJI DZIAŁALNOŚCI (art. 5 ust. 2 UoR).\n"
            f"Opis: {info.get('zagrozenie_opis','')}\n"
            "Opisz wpływ na wycenę aktywów i pasywów w sekcji 1.2.\n"
        )

    context_parts = [
        f"NAZWA: {info.get('nazwa') or company_name}",
        f"FORMA PRAWNA: {info.get('forma_prawna','')}",
        f"SIEDZIBA: {info.get('siedziba','')}",
        f"NIP: {info.get('nip','')}",
        f"KRS: {info.get('krs','')}",
        f"REGON: {info.get('regon','')}",
        f"PKD: {info.get('pkd','')}",
        f"DATA REJESTRACJI: {info.get('data_rejestracji','')}",
        f"OKRES: od {info.get('okres_od','')} do {info.get('okres_do','')}",
        f"ZATRUDNIENIE bieżący rok: {info.get('zatrudnienie_biezacy',0)} etatów",
        f"ZATRUDNIENIE poprzedni rok: {info.get('zatrudnienie_poprzedni',0)} etatów",
        f"UWAGI ZATRUDNIENIE: {info.get('zatrudnienie_uwagi','')}",
        f"ROK OBROTOWY: {year}",
        polityka_blok, zagrozenie_blok,
        "=" * 60,
        "DOKUMENTY FINANSOWE:",
    ]
    for filename, doc_data in doc_mapping.items():
        context_parts.append(f"\n[{doc_data['type']}] {filename}:")
        context_parts.append(doc_data["text"][:8000])
        if len(doc_data["text"]) > 8000:
            context_parts.append("...[skrócono]")

    user_prompt = (
        f"Sporządź Informację Dodatkową za rok {year}.\n\n"
        + "\n".join(context_parts)
        + "\n\nPamiętaj: pomiń noty z zerami, nie pisz ich tytułów. "
        + "Na końcu umieść blok ```wykres_dane z danymi JSON do wykresów."
    )

    if progress_callback:
        progress_callback(0.7, "Generowanie przez Claude...")

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}]
    )
    return response.content[0].text


# ═══════════════════════════════════════════════════════════════════════════════
# MODUŁ 5: WYKRESY SŁUPKOWE
# ═══════════════════════════════════════════════════════════════════════════════

# Paleta kolorów — jednolita w całym dokumencie
BAR_COLORS_PREV = "#A8C4E0"   # jasnoniebieski = rok poprzedni
BAR_COLORS_CURR = "#1B2A4A"   # granat = rok bieżący
BAR_SINGLE      = "#2D6A9F"   # niebieski = pojedynczy rok
GRID_COLOR      = "#E5E5E5"
FONT_NAME       = "DejaVu Sans"


def _fmt_pln(val, _pos=None):
    """Formatuje oś Y w tysiącach PLN."""
    if abs(val) >= 1_000_000:
        return f"{val/1_000_000:.1f} mln"
    if abs(val) >= 1_000:
        return f"{val/1_000:.0f} tys."
    return f"{val:.0f}"


def _parse_chart_data(generated_text: str) -> dict | None:
    """Wyciąga blok JSON wykresów z odpowiedzi Claude."""
    match = re.search(r"```wykres_dane\s*\n(.*?)\n```", generated_text, re.DOTALL)
    if not match:
        return None
    try:
        return json.loads(match.group(1))
    except Exception:
        return None


def _setup_ax(ax, title: str, labels_f, year_suffix: str = ""):
    """Wspólna konfiguracja osi wykresu — styl jednolity w całym dokumencie."""
    ax.set_title(f"{title}{year_suffix}", fontsize=13, fontweight="bold",
                 color="#1B2A4A", pad=12, fontfamily=FONT_NAME)
    ax.set_xticks(np.arange(len(labels_f)))
    ax.set_xticklabels(labels_f, fontsize=9, fontfamily=FONT_NAME)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(_fmt_pln))
    ax.tick_params(axis="y", labelsize=8)
    ax.grid(axis="y", color=GRID_COLOR, linewidth=0.8, zorder=0)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    for spine in ("left", "bottom"):
        ax.spines[spine].set_color(GRID_COLOR)


def _bar_labels(ax, bars, fontsize=7.5):
    """Dodaje etykiety wartości na słupkach."""
    for bar in bars:
        h = bar.get_height()
        if h != 0:
            ax.text(bar.get_x() + bar.get_width()/2, h,
                    _fmt_pln(h), ha="center", va="bottom",
                    fontsize=fontsize, color="#333333", fontfamily=FONT_NAME)


def _save_fig(fig) -> bytes:
    """Zapisuje figurę matplotlib do bytes PNG."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _make_bar_chart(title: str, labels: list, values_prev: list,
                    values_curr: list, year: int):
    """Wykres słupkowy porównawczy (dwa lata)."""
    pairs = [(l, p, c) for l, p, c in zip(labels, values_prev, values_curr)
             if p != 0 or c != 0]
    if not pairs:
        return None
    labels_f, prev_f, curr_f = zip(*pairs)
    x = np.arange(len(labels_f))
    width = 0.38

    fig, ax = plt.subplots(figsize=(10, 4.5), dpi=130)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    bars_p = ax.bar(x - width/2, prev_f, width, color=BAR_COLORS_PREV,
                    label=str(year - 1), zorder=3)
    bars_c = ax.bar(x + width/2, curr_f, width, color=BAR_COLORS_CURR,
                    label=str(year), zorder=3)

    _setup_ax(ax, title, labels_f)
    ax.legend(fontsize=9, framealpha=0)
    _bar_labels(ax, list(bars_p) + list(bars_c), fontsize=7.5)
    plt.tight_layout()
    return _save_fig(fig)


def _make_single_bar_chart(title: str, labels: list, values: list, year: int):
    """Wykres słupkowy dla jednego roku."""
    pairs = [(l, v) for l, v in zip(labels, values) if v != 0]
    if not pairs:
        return None
    labels_f, vals_f = zip(*pairs)

    fig, ax = plt.subplots(figsize=(8, 4), dpi=130)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    bars = ax.bar(np.arange(len(labels_f)), vals_f, color=BAR_SINGLE, zorder=3, width=0.5)
    _setup_ax(ax, title, labels_f, year_suffix=f" ({year})")
    _bar_labels(ax, bars, fontsize=8)
    plt.tight_layout()
    return _save_fig(fig)


def build_charts(chart_data: dict, year: int) -> list[tuple[str, bytes]]:
    """
    Buduje wszystkie wykresy z danych JSON.
    Zwraca listę (tytuł, png_bytes).
    """
    charts = []

    # Wykres 1: Wynik finansowy (porównawczy)
    w = chart_data.get("wynik", {})
    if w:
        png = _make_bar_chart(
            w.get("tytul", "Wynik finansowy"),
            w.get("etykiety", []),
            w.get("rok_poprzedni", []),
            w.get("rok_biezacy", []),
            year
        )
        if png:
            charts.append((w.get("tytul", "Wynik finansowy"), png))

    # Wykres 2: Koszty rodzajowe (porównawczy)
    k = chart_data.get("koszty", {})
    if k:
        png = _make_bar_chart(
            k.get("tytul", "Struktura kosztów"),
            k.get("etykiety", []),
            k.get("rok_poprzedni", []),
            k.get("rok_biezacy", []),
            year
        )
        if png:
            charts.append((k.get("tytul", "Struktura kosztów"), png))

    # Wykres 3: Pasywa (pojedynczy rok)
    p = chart_data.get("pasywa", {})
    if p:
        png = _make_single_bar_chart(
            p.get("tytul", "Struktura pasywów"),
            p.get("etykiety", []),
            p.get("rok_biezacy", []),
            year
        )
        if png:
            charts.append((p.get("tytul", "Struktura pasywów"), png))

    return charts


# ═══════════════════════════════════════════════════════════════════════════════
# MODUŁ 6: EKSPORT DO WORD
# ═══════════════════════════════════════════════════════════════════════════════

NAVY  = RGBColor(0x1B, 0x2A, 0x4A)
BLUE  = RGBColor(0x2D, 0x6A, 0x9F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x22, 0x22, 0x22)
GRAY6 = RGBColor(0x66, 0x66, 0x66)
GRAY9 = RGBColor(0x99, 0x99, 0x99)
LIGHT = "EBF3FB"


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2D6A9F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _tcPr_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _cell_margins(cell, val=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_runs(para, text, bold=False, color=None, size_pt=None):
    run = para.add_run(text)
    run.font.name = "Calibri"
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def _add_inline_text(doc, line, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for part in re.split(r"(\*\*[^*]+\*\*)", line):
        if part.startswith("**") and part.endswith("**"):
            _add_runs(p, part[2:-2], bold=True)
        else:
            _add_runs(p, part, bold=False)
    return p


def _render_md_table(doc, table_lines):
    rows_raw = [l for l in table_lines
                if not re.match(r"^\|[\s\-:|]+\|$", l.strip())]
    if not rows_raw:
        return
    rows = [[c.strip() for c in l.strip().strip("|").split("|")] for l in rows_raw]
    ncols = max(len(r) for r in rows)
    rows  = [r + [""] * (ncols - len(r)) for r in rows]
    col_w = 8500 // ncols

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for ri, row_data in enumerate(rows):
        is_header = (ri == 0)
        for ci, raw_text in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.width = Pt(col_w)
            if is_header:
                _tcPr_shading(cell, "1B2A4A")
            elif ri % 2 == 1:
                _tcPr_shading(cell, LIGHT)
            _cell_margins(cell, 80)

            clean = re.sub(r"[*][*]([^*]+)[*][*]", r"\1", raw_text)
            clean = "".join(ch for ch in clean if ord(ch) >= 32 or ord(ch) in (9, 10, 13))
            is_bold_md = "**" in raw_text

            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(clean)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.bold = is_header or is_bold_md
            run.font.color.rgb = WHITE if is_header else DARK
            if is_header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif re.search(r"\d", clean):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()


def _strip_chart_block(text: str) -> str:
    """Usuwa blok ```wykres_dane z tekstu przed zapisem do Word."""
    return re.sub(r"```wykres_dane.*?```", "", text, flags=re.DOTALL).strip()


def save_to_word(generated_text: str, company_name: str, year: int,
                 charts: list = None) -> bytes:
    doc = Document()

    # Style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    for lvl, sz, col in [(1, 14, NAVY), (2, 12, BLUE), (3, 11, NAVY)]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = "Calibri"
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = col

    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(1.8)

    # Strona tytułowa
    h = doc.add_paragraph()
    _add_runs(h, f"{company_name} | Informacja Dodatkowa {year}",
              bold=False, color=GRAY6, size_pt=9)
    h.paragraph_format.space_after = Pt(12)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_runs(t, "INFORMACJA DODATKOWA", bold=True, color=NAVY, size_pt=22)
    t.paragraph_format.space_before = Pt(24)
    t.paragraph_format.space_after  = Pt(4)

    for txt, sz in [("do sprawozdania finansowego", 14), (f"za rok obrotowy {year}", 14)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_runs(p, txt, bold=False, color=BLUE, size_pt=sz)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(16)

    for label, val in [("Jednostka", company_name),
                        ("Okres sprawozdawczy", f"01.01.{year} — 31.12.{year}")]:
        p = doc.add_paragraph()
        _add_runs(p, f"{label}: ", bold=False, color=GRAY6, size_pt=10)
        _add_runs(p, val, bold=True, color=GRAY6, size_pt=10)

    p_date = doc.add_paragraph()
    _add_runs(p_date, f"Wygenerowano: {datetime.now().strftime('%d.%m.%Y')}",
              bold=False, color=GRAY9, size_pt=9)
    p_date.paragraph_format.space_before = Pt(8)
    doc.add_page_break()

    # Treść — bez bloku wykres_dane
    content = _strip_chart_block(generated_text)
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        strip = lines[i].strip()
        if not strip:
            i += 1
            continue

        if strip.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            _render_md_table(doc, tbl)
            continue

        if strip.startswith("#### "):
            h = doc.add_heading(strip[5:], level=4)
            for r in h.runs: r.font.name = "Calibri"
        elif strip.startswith("### "):
            doc.add_heading(strip[4:], level=3)
        elif strip.startswith("## "):
            doc.add_heading(strip[3:], level=2)
        elif strip.startswith("# "):
            doc.add_heading(strip[2:], level=1)
        elif strip.startswith("---"):
            add_horizontal_rule(doc)
        elif strip.startswith("- ") or strip.startswith("* "):
            _add_inline_text(doc, strip[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", strip):
            _add_inline_text(doc, strip, style="List Number")
        else:
            _add_inline_text(doc, strip)
        i += 1

    # Wykresy
    if charts:
        doc.add_page_break()
        heading = doc.add_heading("Analiza graficzna", level=1)
        for r in heading.runs:
            r.font.color.rgb = NAVY

        for chart_title, png_bytes in charts:
            h3 = doc.add_heading(chart_title, level=3)
            for r in h3.runs:
                r.font.color.rgb = BLUE
            img_stream = io.BytesIO(png_bytes)
            doc.add_picture(img_stream, width=Cm(16))
            doc.add_paragraph()

    # Stopka
    add_horizontal_rule(doc)
    foot = doc.add_paragraph(f"Informacja Dodatkowa | {company_name} | Rok {year}")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in foot.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY9

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _sanitize_text(text: str) -> str:
    import unicodedata
    return "".join(
        ch for ch in text
        if (ord(ch) in (9, 10, 13) or ord(ch) >= 32)
        and unicodedata.category(ch) != "Cs"
    )


# ═══════════════════════════════════════════════════════════════════════════════
# GŁÓWNY INTERFEJS
# ═══════════════════════════════════════════════════════════════════════════════

_anthropic_from_secrets = st.secrets.get("ANTHROPIC_API_KEY", "")
_llama_from_secrets     = st.secrets.get("LLAMA_API_KEY", "")
_app_password           = st.secrets.get("APP_PASSWORD", "")

if _app_password:
    if not st.session_state.get("authenticated"):
        st.markdown("""
        <div style="max-width:400px;margin:4rem auto;padding:2rem;
                    border:1px solid #dee2e6;border-radius:12px;
                    box-shadow:0 4px 12px rgba(0,0,0,0.1);text-align:center;">
            <h2>🔐 Dostęp chroniony</h2>
            <p style="color:#666;">Wprowadź hasło aby kontynuować</p>
        </div>""", unsafe_allow_html=True)
        _, col_mid, _ = st.columns([1, 2, 1])
        with col_mid:
            entered = st.text_input("Hasło", type="password",
                                    label_visibility="collapsed",
                                    placeholder="Wpisz hasło...")
            if st.button("Zaloguj →", use_container_width=True, type="primary"):
                if entered == _app_password:
                    st.session_state["authenticated"] = True
                    st.rerun()
                else:
                    st.error("❌ Nieprawidłowe hasło")
        st.stop()

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Konfiguracja")

    if _anthropic_from_secrets:
        st.success("🔑 Klucz API Anthropic: wczytany automatycznie")
        anthropic_key = _anthropic_from_secrets
    else:
        anthropic_key = st.text_input("🔑 Klucz API Anthropic", type="password",
                                       placeholder="sk-ant-...")

    if _llama_from_secrets:
        st.success("🦙 Klucz LlamaParse: wczytany automatycznie")
        llama_key = _llama_from_secrets
    else:
        llama_key = st.text_input("🦙 Klucz LlamaParse (opcjonalny)", type="password",
                                   placeholder="llx-...")

    st.divider()
    st.subheader("🏢 Dane jednostki")

    krs_input = st.text_input("🔍 Numer KRS", placeholder="0000123456",
                               help="10-cyfrowy numer KRS — znajdziesz na prs.ms.gov.pl")
    st.caption("ℹ️ API KRS działa po numerze KRS (nie NIP).")
    debug_krs = st.checkbox("🔍 Tryb diagnostyczny KRS", value=False)

    if st.button("⬇️ Pobierz dane z KRS", use_container_width=True):
        if krs_input:
            with st.spinner("Pobieranie z API KRS..."):
                try:
                    if debug_krs:
                        krs_data, log = fetch_krs_by_krs_nr_debug(krs_input)
                        st.code(log)
                    else:
                        krs_data = fetch_krs_by_krs_nr(krs_input)
                    if krs_data:
                        st.session_state["krs_data"] = krs_data
                        st.success("✅ Dane pobrane z KRS!")
                    else:
                        st.error("❌ Nie znaleziono. Sprawdź numer KRS.")
                except Exception as e:
                    st.error(f"❌ {e}")
        else:
            st.warning("Wpisz numer KRS.")

    krs = st.session_state.get("krs_data", {})

    company_name     = st.text_input("Nazwa spółki",       value=krs.get("nazwa", ""),          placeholder="XYZ Sp. z o.o.")
    company_siedziba = st.text_input("Siedziba",           value=krs.get("siedziba", ""),        placeholder="ul. Przykładowa 1, 00-001 Warszawa")
    company_nip      = st.text_input("NIP",                value=krs.get("nip", ""),             placeholder="1234567890")
    company_krs      = st.text_input("Nr KRS",             value=krs.get("krs", krs_input or ""),placeholder="0000000000")
    company_regon    = st.text_input("REGON",              value=krs.get("regon", ""),           placeholder="000000000")
    company_pkd      = st.text_input("Główne PKD",         value=krs.get("pkd", ""),             placeholder="69.20.Z Rachunkowość")
    company_data_rej = st.text_input("Data rejestracji",   value=krs.get("data_rejestracji",""), placeholder="DD.MM.RRRR")
    company_forma    = st.text_input("Forma prawna",       value=krs.get("forma_prawna", ""),    placeholder="SPÓŁKA Z OGRANICZONĄ ODPOWIEDZIALNOŚCIĄ")

    st.divider()
    st.subheader("📅 Okres sprawozdawczy")
    okres_od   = st.date_input("Od", value=date(date.today().year - 1, 1,  1))
    okres_do   = st.date_input("Do", value=date(date.today().year - 1, 12, 31))
    fiscal_year = okres_do.year

    st.divider()
    st.subheader("⚠️ Kontynuacja działalności")
    zagrozenie_kontynuacji = st.checkbox(
        "Istnieją okoliczności wskazujące na zagrożenie kontynuowania działalności "
        "w okresie co najmniej 12 miesięcy od dnia bilansowego",
        value=False, help="Art. 5 ust. 2 UoR"
    )
    zagrozenie_opis = ""
    if zagrozenie_kontynuacji:
        zagrozenie_opis = st.text_area("Opis okoliczności:", height=100)

    st.divider()
    st.subheader("👥 Zatrudnienie")
    zatrudnienie_biezacy   = st.number_input("Rok bieżący (etaty)",  min_value=0, value=0, step=1)
    zatrudnienie_poprzedni = st.number_input("Rok poprzedni (etaty)",min_value=0, value=0, step=1)
    zatrudnienie_uwagi     = st.text_input("Uwagi", placeholder="np. w tym 2 osoby na zleceniu")

    st.divider()
    st.markdown("""
    **📋 Obsługiwane dokumenty:**
    - 🏦 Bilans
    - 📈 Rachunek Zysków i Strat
    - 🏗️ Tabela środków trwałych
    - 💸 Przepływy pieniężne
    - 📜 Polityka rachunkowości
    - ⚖️ Zestawienie Obrotów i Sald
    """)

# ── GŁÓWNA SEKCJA ─────────────────────────────────────────────────────────────
col1, col2 = st.columns(2)
with col1:
    st.markdown('<div class="step-card"><b>📁 Krok 1:</b> Wgraj dokumenty PDF</div>',
                unsafe_allow_html=True)
    uploaded_files = st.file_uploader("Wybierz pliki PDF", type=["pdf"],
                                       accept_multiple_files=True)
    if uploaded_files:
        st.success(f"✅ Wgrano {len(uploaded_files)} plik(ów)")
        for f in uploaded_files:
            st.caption(f"📄 {f.name} ({len(f.getvalue())//1024} KB)")

with col2:
    st.markdown('<div class="step-card"><b>🔍 Krok 2:</b> Walidacja i generowanie</div>',
                unsafe_allow_html=True)
    if not anthropic_key:
        st.info("👈 Wprowadź klucz API Anthropic w panelu bocznym.")
    elif not uploaded_files:
        st.info("👈 Wgraj pliki PDF.")
    elif not company_name:
        st.warning("⚠️ Wprowadź nazwę spółki.")

# ═══════════════════════════════════════════════════════════════════════════════
# MASZYNA STANÓW
# ═══════════════════════════════════════════════════════════════════════════════

st.divider()

def _reset_state():
    for k in ["app_state","parsed_docs","doc_mapping","missing_docs",
              "polityka_answers","generated_text","docx_bytes"]:
        st.session_state.pop(k, None)

def _set_state(s): st.session_state["app_state"] = s
def _get_state(): return st.session_state.get("app_state", "idle")

run_disabled = not (anthropic_key and uploaded_files and company_name)
if st.button("🚀 Generuj Informację Dodatkową", type="primary",
             disabled=run_disabled, use_container_width=True):
    _reset_state()
    _set_state("parsing")
    st.rerun()

# ── STAN: parsing ─────────────────────────────────────────────────────────────
if _get_state() == "parsing":
    pb = st.progress(0); st_txt = st.empty()
    try:
        st_txt.info("📄 Parsowanie dokumentów PDF...")
        pb.progress(10)

        def _upd(v, m): pb.progress(int(10 + v * 20)); st_txt.info(f"📄 {m}")

        parsed = (parse_documents_llamaparse(uploaded_files, llama_key, _upd)
                  if llama_key else parse_documents_fallback(uploaded_files, _upd))

        st_txt.info("🗂️ Mapowanie dokumentów...")
        pb.progress(40)
        doc_mapping = map_documents(parsed)

        st.session_state["parsed_docs"]  = parsed
        st.session_state["doc_mapping"]  = doc_mapping
        st.session_state["missing_docs"] = check_missing_documents(doc_mapping)

        pb.empty(); st_txt.empty()
        _set_state("confirm_missing" if st.session_state["missing_docs"] else "polityka")
        st.rerun()
    except Exception as e:
        pb.empty(); st_txt.empty()
        _set_state("error"); st.session_state["error_msg"] = str(e); st.rerun()

# ── STAN: confirm_missing ─────────────────────────────────────────────────────
elif _get_state() == "confirm_missing":
    missing    = st.session_state.get("missing_docs", [])
    doc_mapping = st.session_state.get("doc_mapping", {})

    st.subheader("📋 Rozpoznane dokumenty")
    cols = st.columns(max(len(doc_mapping), 1))
    for i, (fname, ddata) in enumerate(doc_mapping.items()):
        with cols[i % len(cols)]:
            st.markdown(f'<div class="metric-box"><b>{ddata["type"]}</b><br>'
                        f'<small>{fname}</small></div>', unsafe_allow_html=True)

    st.warning("⚠️ Nie znaleziono wszystkich dokumentów.")
    st.markdown("**Brakujące:**")
    for dt in missing:
        info_dt = REQUIRED_DOC_TYPES[dt]
        st.markdown(f"- {info_dt['icon']} **{info_dt['label']}** — {info_dt['desc']}")

    st.info("💡 Jeśli kilka dokumentów jest w jednym PDF — spróbuj je rozdzielić.")
    st.markdown("---")
    col_a, col_b = st.columns(2)
    with col_a:
        if st.button("▶️ Kontynuuj bez brakujących",
                     use_container_width=True, type="primary"):
            _set_state("polityka"); st.rerun()
    with col_b:
        if st.button("📁 Anuluj — chcę dodać pliki", use_container_width=True):
            _reset_state(); st.rerun()

# ── STAN: polityka ────────────────────────────────────────────────────────────
elif _get_state() == "polityka":
    doc_mapping = st.session_state.get("doc_mapping", {})
    types_found = {d["type"] for d in doc_mapping.values()}

    if "POLITYKA RACHUNKOWOŚCI" not in types_found:
        st.warning("📜 Brak Polityki Rachunkowości — wypełnij poniższe pytania.")
        with st.form("polityka_form"):
            st.subheader("📋 Zasady rachunkowości")
            q1 = st.selectbox("1. Zasady ustalania wyniku:", [
                "Wariant porównawczy (układ rodzajowy kosztów)",
                "Wariant kalkulacyjny (układ funkcjonalny kosztów)"])
            q2a = st.selectbox("2a. Wycena zapasów:", [
                "FIFO", "LIFO", "Cena przeciętna (średnia ważona)",
                "Ceny ewidencyjne z odchyleniami", "Nie dotyczy (brak zapasów)"])
            q2b = st.selectbox("2b. Amortyzacja ST:", [
                "Liniowa", "Degresywna",
                "Jednorazowy odpis (niskocenne do 10 000 zł)", "Mieszana"])
            q2c = st.selectbox("2c. Wycena należności:", [
                "Wartość nominalna z odpisami aktualizującymi",
                "Wartość nominalna bez odpisów", "Wartość godziwa"])
            q3 = st.selectbox("3. Rodzaj sprawozdania:", [
                "Pełne sprawozdanie finansowe",
                "Uproszczone (art. 46 ust. 5 UoR — jednostki małe)",
                "Załącznik nr 4 UoR (mikro jednostki)",
                "Załącznik nr 5 UoR (małe NGO)"])
            q4 = st.checkbox("Tworzy rezerwę/aktywa z tytułu CIT odroczonego", value=True)
            q5 = st.selectbox("Leasing:", [
                "Wg UoR (operacyjny/finansowy wg treści ekonomicznej)",
                "Tylko operacyjny",
                "Nie dotyczy (brak umów leasingowych)"])
            uwagi = st.text_area("Uwagi dodatkowe:", height=70)

            if st.form_submit_button("✅ Zatwierdź i generuj",
                                     use_container_width=True, type="primary"):
                st.session_state["polityka_answers"] = {
                    "wynik_finansowy": q1, "wycena_zapasow": q2a,
                    "amortyzacja": q2b, "wycena_naleznosci": q2c,
                    "sposob_sprawozdania": q3, "podatek_odroczony": q4,
                    "leasing": q5, "uwagi": uwagi,
                }
                _set_state("generating"); st.rerun()
    else:
        st.session_state["polityka_answers"] = {}
        _set_state("generating"); st.rerun()

# ── STAN: generating ──────────────────────────────────────────────────────────
elif _get_state() == "generating":
    doc_mapping      = st.session_state.get("doc_mapping", {})
    polityka_answers = st.session_state.get("polityka_answers", {})

    pb = st.progress(0); st_txt = st.empty(); results_box = st.container()

    try:
        # Walidacja
        st_txt.info("✅ Walidacja spójności danych...")
        pb.progress(55)
        issues = validate_data_consistency(doc_mapping)

        with results_box:
            st.subheader("📋 Rozpoznane dokumenty")
            cols = st.columns(max(len(doc_mapping), 1))
            for i, (fname, ddata) in enumerate(doc_mapping.items()):
                with cols[i % len(cols)]:
                    st.markdown(f'<div class="metric-box"><b>{ddata["type"]}</b><br>'
                                f'<small>{fname}</small><br>'
                                f'<small>{ddata["length"]:,} znaków</small></div>',
                                unsafe_allow_html=True)
            st.subheader("🔎 Walidacja")
            css = {"OK":"validation-ok","WARN":"validation-warn","ERR":"validation-err"}
            for issue in issues:
                st.markdown(f'<span class="{css.get(issue["level"],"")}">{ issue["msg"]}</span>',
                            unsafe_allow_html=True)

        # Generowanie
        st_txt.info("🤖 Generowanie przez Claude 3.5 Sonnet (1–3 min)...")
        pb.progress(65)

        company_info = {
            "nazwa": company_name,         "siedziba": company_siedziba,
            "nip": company_nip,             "krs": company_krs,
            "regon": company_regon,         "pkd": company_pkd,
            "data_rejestracji": company_data_rej, "forma_prawna": company_forma,
            "okres_od": str(okres_od),      "okres_do": str(okres_do),
            "zagrozenie_kontynuacji": zagrozenie_kontynuacji,
            "zagrozenie_opis": zagrozenie_opis,
            "polityka_answers": polityka_answers,
            "zatrudnienie_biezacy": zatrudnienie_biezacy,
            "zatrudnienie_poprzedni": zatrudnienie_poprzedni,
            "zatrudnienie_uwagi": zatrudnienie_uwagi,
        }

        generated_text = generate_accounting_notes(
            doc_mapping=doc_mapping,
            anthropic_api_key=anthropic_key,
            company_name=company_name,
            year=fiscal_year,
            company_info=company_info,
            progress_callback=lambda v, m: pb.progress(int(65 + v * 15))
        )
        pb.progress(82)

        # Wykresy
        st_txt.info("📊 Generowanie wykresów...")
        clean_text  = _sanitize_text(generated_text)
        chart_data  = _parse_chart_data(clean_text)
        charts      = build_charts(chart_data, fiscal_year) if chart_data else []
        pb.progress(92)

        # Word
        st_txt.info("💾 Zapis do Word...")
        try:
            docx_bytes = save_to_word(clean_text, company_name, fiscal_year, charts)
        except Exception as docx_err:
            import traceback
            st.error(f"❌ Błąd Word: {docx_err}")
            st.code(traceback.format_exc())
            st.stop()

        st.session_state["generated_text"] = generated_text
        st.session_state["docx_bytes"]     = docx_bytes
        pb.progress(100)
        st_txt.success("✅ Informacja Dodatkowa wygenerowana!")
        _set_state("done"); st.rerun()

    except anthropic.AuthenticationError:
        _set_state("error"); st.session_state["error_msg"] = "Nieprawidłowy klucz API."
        st.rerun()
    except anthropic.RateLimitError:
        _set_state("error"); st.session_state["error_msg"] = "Limit API — poczekaj chwilę."
        st.rerun()
    except Exception as e:
        _set_state("error"); st.session_state["error_msg"] = str(e); st.rerun()

# ── STAN: done ────────────────────────────────────────────────────────────────
elif _get_state() == "done":
    st.success("✅ Informacja Dodatkowa wygenerowana!")
    col_dl, _ = st.columns([1, 2])
    with col_dl:
        st.download_button(
            "⬇️ Pobierz (.docx)",
            data=st.session_state["docx_bytes"],
            file_name=f"informacja_dodatkowa_{company_name.replace(' ','_')}_{fiscal_year}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary", use_container_width=True
        )
    if st.button("🔄 Generuj dla innej spółki", use_container_width=True):
        _reset_state(); st.rerun()
    with st.expander("👁️ Podgląd treści", expanded=False):
        st.markdown(st.session_state["generated_text"])

# ── STAN: error ───────────────────────────────────────────────────────────────
elif _get_state() == "error":
    st.error(f"❌ {st.session_state.get('error_msg','Nieznany błąd')}")
    if st.button("🔄 Spróbuj ponownie"):
        _reset_state(); st.rerun()
